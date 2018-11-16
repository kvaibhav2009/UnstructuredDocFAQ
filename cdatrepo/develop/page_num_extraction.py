from tika import parser
from pandas import DataFrame
from bs4 import BeautifulSoup
#import Doc2PDF_pg_num as doc2pdf
## Filepath to the pdf
import pandas as pd
import string
import sys
import os
import re
from docx import Document
from nltk.util import ngrams
import json
from StringIO import StringIO
import pycurl
import certifi

reload(sys)
sys.setdefaultencoding('utf8')
# 8 /29/2018
# pdf_file=doc2pdf.Doc2PDFConverter(inputFile)
#
# #pdf_file = r'C:\Users\binit.kumar.bhagat\Documents\CDAT\Page_no_prob\dumps\Liability -  AI sample 1 .pdf'
#
# pdf_parsed = parser.from_file(pdf_file, xmlContent=True)
#
# soup = BeautifulSoup(pdf_parsed['content'], "lxml")
#
# text_list = []
# counter = 1
# for page in soup.find_all("div"):
#     print('\n\n'+'Page '+str(counter))
#     for sect in page.find_all("p"):
#         print(sect.text)
#         text_dict = {"sect": sect.text, "page": counter}
#         text_list.append(text_dict)
#     counter += 1
# f='C:\\Enclave\\Git projects\\Text Analytics\\CDAT_Flask\\Input Documents\\Benchmarking -  AI sample 1.docx'
# word=['benchmarking', 'benchmark']

paraindex=[]

def createSectionFile(f,docxpath,words):
    #f="C:\Enclave\Git projects\Text Analytics\CDAT Flask\outputExcel\sample 5.pdf"
    #f="FileConversion/sample 5.pdf"
    print("Creating section file for "+str(words))
    print("File")
    text_list=textlist(f)
    text_listdf = pd.DataFrame(text_list)
    text_listdf.to_csv("text_listdf")
    # f.save((f.filename))
    #doc = Document(f)
    print("Document")
    docxpath=docxpath #"C:\Enclave\Git projects\Text Analytics\CDAT Flask\FileConversion\sample 5doc.docx"
    para=[]
    doc=Document(docxpath)
    if('liability' in words):
        for i,paragraphs in enumerate(doc.paragraphs):
                if (str(paragraphs.text.lower()).__contains__('limitation') & str(paragraphs.text.lower()).__contains__('liability')) :#(str(paragraphs.text.lower()).__contains__('limitation of liability')):   #x
                    para.append(paragraphs.text)
                    if (str(doc.paragraphs[i + 1].text)):

                        text=''.join(x for x in (str(doc.paragraphs[i + 1].text).encode('utf-8').decode('latin-1')) if x in string.printable)
                        para.append(text)
                    else:
                        text=''.join(x for x in (str(doc.paragraphs[i + 2].text).encode('utf-8').decode('latin-1')) if x in string.printable)
                        para.append(text)

    if (para.__len__() == 0):
        print("Specific section limitation of liability not found")
        for i, paragraphs in enumerate(doc.paragraphs):
            for x in words:
                if (str(paragraphs.text.lower()).__contains__(x)):  # x
                        para.append(paragraphs.text)
    # Commment if doesnot work
    if('benchmarking' not in words):

        para,xpara= constructpara(docxpath, words)
        para=xpara
    # Commment if doesnot work

    setpara=set(para)
    paradataframe = pd.DataFrame(list(setpara))
    paradataframe.to_csv("Doc_extracted_para.csv")
    para=list(setpara)


    print("Para list length: "+str(para.__len__()))

    l = list()
    section = list()
    file = "Output Documents\\" + "sectioning" + ".csv"
    file=os.path.join("Output Documents","sectioning.csv")
    sectionfile = "OutputSection\\" + "sectioning" + "_section.csv"
    sectionfile=os.path.join("OutputSection","sectioning_section.csv")
    indexfile=list()
    print("Creating the sections: ")
    parapdfdataframe = pd.DataFrame((text_list))
    parapdfdataframe.to_csv("PDF_extracted_para.csv")
    # ///////////////////////////////
    # for i, paragraphs in enumerate(text_list):
    #     text = (paragraphs['sect'])
    #     text = ''.join(x for x in (text.encode('utf-8').decode('latin-1')) if x in string.printable)
    #
    #     if str(text):
    #         l.append(str(text))
    #     for x in words:
    #         if (str(text.lower()).__contains__(
    #                 x)):  # if [s for s in (text.lower()).split() if any(xs in s for xs in words)]:
    #             #text = preprocess(text)
    #             print("Testing for given section of pdf "+ str(text.lower()))
    #             paraText= extractParagraph(para,str(text.lower()))
    #             print("ParaText: "+str(paraText))
    #             print("If paratext empty "+str(not str(paraText)))
    #             if(bool(str(paraText))):
    #                 text = ''.join(x for x in str(paraText) if x in string.printable)
    #                 section.append(text)
    #                 index=[(text),paragraphs['page'],paragraphs['sect_no']]
    #                 print("\n Index record " + str(index))
    #                 indexfile.append(index)
    #                 #para.remove(str(paraText))
    #             break;
    # d = pd.DataFrame(l)
    # sectiondf = pd.DataFrame(section)
    # sectiondf.to_csv(sectionfile)
    # d.to_csv(file)

    pdflist=parapdfdataframe['sect'].tolist()

    for docpara in para:
        match=[]
        docparatext = ''.join(x for x in (docpara.encode('utf-8').decode('latin-1')) if x in string.printable)
        section.append(str(docparatext))
        for sectiontext in pdflist:
            if (str(sectiontext)):
                match.append(int(String_Subsequence(str(docpara), str(sectiontext))))
            else:
                match.append(0)
        indexno=int(match.index(max(match)))
        if(indexno>0):
            print(pdflist[indexno])
        parapdfdataframe['sect_no'][indexno]
        if(str(pdflist[indexno]).__len__()>docpara.__len__()):
            docparatext = ''.join(x for x in (str(pdflist[indexno]).encode('utf-8').decode('latin-1')) if x in string.printable)
        else:
            docparatext=''.join(x for x in (docpara.encode('utf-8').decode('latin-1')) if x in string.printable)

        index = [str(docparatext), parapdfdataframe['page'][indexno], parapdfdataframe['sect_no'][indexno]]
        indexfile.append(index)
        section.append(str(docparatext))
    # /////////////////////////////////////
    section=list(set(section))
    sectiondf = pd.DataFrame(section)
    #sectiondf=sectiondf.drop_duplicates(subset=[sectiondf.columns[0]])
    sectiondf.to_csv(sectionfile,index=False)

    indexfiledf = pd.DataFrame(indexfile)
    #indexfiledf = indexfiledf.drop_duplicates(subset=[indexfiledf.columns[0]])
    indexfilename="OutputSection\\" + "sectioning_dict" + ".csv"
    indexfilename=os.path.join("OutputSection","sectioning_dict.csv")
    indexfiledf.to_csv(indexfilename,index=False)

    return sectionfile,indexfilename

def preprocess(text):
    dictword = {"twice": "200%", "thrice": "300%","exceeds":"200%"}  #,"exceed":"200%","exceeds":"200%","exceeded":"200%"
    for word in text.split():
        if (word in dictword.keys()):
            text=text.replace(word, dictword[word])

    return text

def textlist(f):
    inputFile=f
    print("inputFile "+inputFile)
    pdf_file = f #doc2pdf.Doc2PDFConverter(inputFile)

    # pdf_file = r'C:\Users\binit.kumar.bhagat\Documents\CDAT\Page_no_prob\dumps\Liability -  AI sample 1 .pdf'

    pdf_parsed = parser.from_file(pdf_file, xmlContent=True)

    soup = BeautifulSoup(pdf_parsed['content'], "lxml")

    text_list = []
    counter = 1
    sect_no = ''
    pattern = '(\d+[.])+\d*\s+'
    print('\n\n' + 'First Page ' + str(counter))
    for page in soup.find_all("div"):
        # print('\n\n' + 'Page ' + str(counter))
        for sect in page.find_all("p"):
            #print(sect.text)
            sector = re.match(pattern, sect.text)
            if sector:
                sect_no = sector.group(0).strip()
            text_dict = {"sect": sect.text, "page": counter,"sect_no": sect_no}
            text_list.append(text_dict)
        counter += 1
    print('\n\n' + 'Last Page ' + str(counter-1))


    #print(text_list)
    return text_list





def longest_common_subsequence(X, Y):
    # find the length of the strings
    m = len(X)
    n = len(Y)

    # declaring the array for storing the dp values
    L = [[None] * (n + 1) for i in xrange(m + 1)]

    """Following steps build L[m+1][n+1] in bottom up fashion 
    Note: L[i][j] contains length of LCS of X[0..i-1] 
    and Y[0..j-1]"""
    for i in range(m + 1):
        for j in range(n + 1):
            if i == 0 or j == 0:
                L[i][j] = 0
            elif X[i - 1] == Y[j - 1]:
                L[i][j] = L[i - 1][j - 1] + 1
            else:
                L[i][j] = max(L[i - 1][j], L[i][j - 1])

                # L[m][n] contains the length of LCS of X[0..n-1] & Y[0..m-1]
    return L[m][n]


def extractParagraph(para,text):
    match = []
    for p in para:
        match.append(int(longest_common_subsequence(p, text)))#(float(String_Subsequence(p,text))) #int(longest_common_subsequence(p, text)))
    print("Match "+str(match))
    try:
        index = match.index(max(match))
    except:
        paraText=''
        return paraText

    if(index in paraindex):
        print("Index " + str(index) + " found repetitve")
        print(str(paraindex))
        paraText = ''
    if(index not in paraindex):
        ("Extracting para at " + str(index))
        paraText = para[index]
        paraindex.append(index)

    return paraText


def String_Subsequence(string1,string2):

    n = 2
    ngrams1 = list(ngrams(string1.split(" "), n))
    ngrams2 = list(ngrams(string2.split(" "), n))

    common_bigrams = [x for x in ngrams1 if x in ngrams2]
    score_bigram = len(common_bigrams)
    #print score_bigram
    #print 'Common Bigrams\n', common_bigrams

    n = 3
    ngrams3 = list(ngrams(string1.split(" "), n))
    ngrams4 = list(ngrams(string2.split(" "), n))

    common_trigrams = [x for x in ngrams3 if x in ngrams4]
    score_trigram = len(common_trigrams)
    #print score_trigram
    #print 'Common Trigrams\n', common_trigrams

    n = 4
    ngrams5 = list(ngrams(string1.split(" "), n))
    ngrams6 = list(ngrams(string2.split(" "), n))

    common_4grams = [x for x in ngrams5 if x in ngrams6]
    score_4gram = len(common_4grams)
    #print score_4gram
    #print 'Common 4-grams\n', common_4grams

    #print '\nTotal score :', 2 * score_bigram + 3 * score_trigram + 4 * score_4gram
    score=float(2 * score_bigram + 3 * score_trigram + 4 * score_4gram)
    return score


#sectionfile,indexfilename=createSectionFile(f,word)
#
# print("Section file "+sectionfile)
# print("Index file "+indexfilename)


def constructpara(docxpath,words):
    print("Constructing para")
    doc = Document(docxpath)
    para = []
    for i, paragraphs in enumerate(doc.paragraphs):
        if (paragraphs.text):
            dtext=''.join(x for x in (str((paragraphs.text)).encode('utf-8').decode('latin-1')) if x in string.printable)
            para.append(dtext)

    sectionfile="Extracted_para.csv"
    target='liability'
    if 'liability' in words:
        pipelineId="1079"
        threshold=0.76
    else:
        target='benchmarking'
        pipelineId = '1204'#"1096"
        threshold = 0.52
    import pandas as pd

    paradataframe = pd.DataFrame(list(para))
    paradataframe.to_csv(sectionfile, index = False)

    data = AIEngineIntegrator(pipelineId, sectionfile)
    count=0
    xpara=[]
    for i in range(json.loads(data['finalOutPut']['finalString'])['Output'].__len__()):
        if (str(json.loads(data['finalOutPut']['finalString'])['Output'][i]['Target']) == target and float(
                json.loads(data['finalOutPut']['finalString'])['Output'][i]['Confidence']) > float(threshold)):
            xpara.append(str(json.loads(data['finalOutPut']['finalString'])['Output'][i]['inputtext']))
            count = count + 1

    return para,xpara


def AIEngineIntegrator(pipelineId,sectionfile):

    storage = StringIO()

    url = "https://aaie.accenture.com/aceapi/ace-pipeline-execution-service/execute/PipelineExecution/"
    filepath = sectionfile#'OutputSection\\liability_section.csv'
    c = pycurl.Curl()
    c.setopt(c.URL, url)
    # c.setopt(c.HTTPHEADER, ['enctype : multipart/form-data', 'Content-Type: multipart/form-data'])

    c.setopt(c.USERPWD, '%s:%s' % ('ace', 'ace'))

    c.setopt(c.CAINFO, certifi.where())

    c.setopt(c.HTTPPOST, [('file', (c.FORM_FILE, filepath,
                                    c.FORM_CONTENTTYPE, 'multipart/form-data',
                                    )), ('ownerUserName', 's.b.karuppasamy'), ('ownerUserId', '130'),
                          ('executorUserName', 's.b.karuppasamy'), ('executorUserId', '130'), ('pipelineId',pipelineId ),
                          ('getStatusWithResponse', 'true'), ('input', ''), ('bulletHandler', 'false'),
                          ('extractText', 'true')

                          ])
    # writing output to storage NER '1041'
    c.setopt(c.WRITEFUNCTION, storage.write)
    c.perform()
    c.close()
    # Required content
    content = storage.getvalue()
    print content
    content_str = content.decode()
    data = json.loads(content_str)

    return data
