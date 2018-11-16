from PyPDF2.utils import readNonWhitespace
from flask import Flask,jsonify
from flask import *
#from werkzeug import secure_filename
from docx import Document
import pandas as pd
import string
import os
import requests
import pycurl
from StringIO import StringIO
import certifi
import Utility.Word2numParser as wp
import re
import Processor as p
from flask import render_template
import sys
from flask_cors import CORS, cross_origin
import page_num_extraction as pg_num
import time
sys.path.append(os.path.abspath(os.path.join(os.path.dirname( __file__ ), '../..')))
import os
#import Doc2PDF_pg_num as doc2pdf
import Doc2PDF_abi_word as abi_doc2pdf
# initializing a variable of Flask
#app = Flask(__name__)
app = Flask(__name__, static_url_path='',template_folder='templates')
import os
app.config['CORS_HEADERS'] = 'CONTENT-TYPE'
cors = CORS(app, resources={r"/*": {"origins": "localhost"}},headers="Content-Type")


# decorating index function with the app.route
@app.route('/')
def index():
   print("Index")
   return render_template("index95.html")

#
# @app.route('/templates/')
# def root():
#     return app.send_static_file('index1.html')

def options(self, *args, **kwargs):
    self.response.headers['Access-Control-Allow-Origin'] = '*'
    self.response.headers[
        'Access-Control-Allow-Headers'] = 'Origin, X-Requested-With, Content-Type, Accept, Authorization'
    self.response.headers['Access-Control-Allow-Methods'] = 'POST, GET, PUT, DELETE'


@app.route('/router', methods=['GET','POST','OPTIONS'])
def router():
    print("Router")
    word = 'liability'
    words = ['liability', 'liable']
    words=['benchmarking','benchmark']
    print(request.form)
    f = request.files['file']
    print("Printing checkboxes")
    print(request.form.getlist('keys[]'))
    print(request.form.getlist('downloadfile'))

    print("File")
    # f.save((f.filename))
    #
    print("Document")
    l = list()
    section = list()
    CheckboxValue=request.form.getlist('keys[]')
    print("CheckboxValue "+str(CheckboxValue))
    print("Type: "+str(type(CheckboxValue)))
    json_response = {'Path':'','Data': []}

    type(f)

    filename = os.getcwd() + "\\OutputFile\\" + f.filename
    filename=os.path.join(os.getcwd(),"OutputFile",f.filename)
    print("Saving " + filename)
    f.save((filename))
    print("Filename " + filename)

    filename=abi_doc2pdf.doc2Pdf(filename)
    for word in CheckboxValue:
        print("Word is "+word)
        print("Boolean: "+str(word is 'Benchmarking'))
        if(str(word)=='Benchmarking'):
            print("It is benchmarking")
            words = ['benchmarking', 'benchmark']
            ClassifierPipelineID='1096'
            NERPipelineID = '1097'
        if(str(word)=='Liability'):
            print("It is liability")
            words=['liability', 'liable']
            ClassifierPipelineID = '1079'
            NERPipelineID='1041'

        sectionFile,indexfilename=pg_num.createSectionFile(filename,"", words)
        print("Section filename "+sectionFile)
        print("Indexfilename "+indexfilename)
        indexfileAbspath=os.getcwd()+"\\"+indexfilename
        indexfileAbspath=os.path.join(os.getcwd(),indexfilename)
        page_nos=get_pagenumbers(indexfileAbspath)
        print("Page Numbers "+page_nos)
        start = time.time()
        ClassifierResponse = AIEngineIntegrator(ClassifierPipelineID, sectionFile)
        done = time.time()
        elapsed = done - start
        print("Time elapsed by classifier"+str(elapsed))
        print("Classifier response", ClassifierResponse)

        start = time.time()
        NERintegration = AIEngineIntegrator(NERPipelineID, sectionFile)
        done = time.time()
        elapsed = done - start
        print("Time elapsed by NER" + str(elapsed))
        print("NERintegration response", NERintegration)

        if (str(word) == 'Liability'):
            DurationData = formulateResponseforDuration(ClassifierResponse, NERintegration)
            AmountData = formulateResponseforAmount(ClassifierResponse, NERintegration)
            outcome, reason = formulateReason(DurationData, AmountData)
            Data = {
                "Checkvalue": "Liability",
                'Confidence': DurationData['Confidence'],
                'Start_End': DurationData['Start_End'],
                'Gap_outcome': outcome,
                'Response': DurationData['Response'],
                'Reason': reason,
                'page_nos':page_nos
            }
            print(Data)
            json_response['Data'].append(Data)
        if (str(word) == 'Benchmarking'):
            # import Utility.NERanalyzer as nr
            #
            #
            # DurationData = formulateResponseforDuration(ClassifierResponse, NERintegration)
            # AmountData = formulateResponseforAmount(ClassifierResponse, NERintegration)
            # outcome, reason = formulateReason(DurationData, AmountData)
            # Data = {
            #     'Confidence': DurationData['Confidence'],
            #     'Start_End': DurationData['Start_End'],
            #     'Gap_outcome': outcome,
            #     'Response': DurationData['Response'],
            #     'Reason': reason
            # }
            # print(Data)
            # json_response['Data'].append(Data)
            notice_period_duration,notice_period_start_end=get_notice_period(sectionFile)
            initial_period_duration,initial_period_start_end = get_initial_period(sectionFile)
            print("Notice_period "+str(notice_period_duration))
            print("Initial_period " +str(initial_period_duration))
            print("Notice period"+str(notice_period_start_end))
            print("Initial period"+str(initial_period_start_end))

            if(notice_period_start_end!=initial_period_start_end):
                start_end="1."+str(notice_period_start_end)+" 2."+str(initial_period_start_end)
            else:
                start_end=str(notice_period_start_end)

            NoticePeriodData=formulateResponseforNoticePeriodBechmarking(notice_period_duration)
            InitialPeriodData = formulateResponseforInitialPeriodBechmarking(initial_period_duration)
            outcome, reason = formulateReasonforBenchmarking(NoticePeriodData,InitialPeriodData)
            Data = {
                "Checkvalue": "Benchmarking",
                'Confidence': 'Confidence_output',
                'Start_End': start_end,
                'Gap_outcome': outcome,
                'Response': NoticePeriodData['Response'],
                'Reason': reason,
                'page_nos': page_nos
            }
            # Data = {
            #     "Checkvalue": "Benchmarking",
            #     'Confidence': 'Confidence_output',
            #     'Start_End': "start_end",
            #     'Gap_outcome': "outcome",
            #     'Response': 'Response',
            #     'Reason': 'reason'
            # }
            json_response['Data'].append(Data)
        #print("Deleting file"+filename)
        clearDirectory()
        print("Directory cleaned")
    # Data = {
    #     "Checkvalue": "Liability",
    #     'Confidence': 'Confidence_output',
    #     'Start_End': 'Start_End_Output',
    #     'Gap_outcome': 'Outcome_output',
    #     'Response': 'Response_output',
    #     'Reason': 'Reason_output'
    # }
    #
    # Data1 = {
    #     "Checkvalue":"Benchmarking",
    #     'Confidence': 'Confidence_output1',
    #     'Start_End': 'Start_End_Output1',
    #     'Gap_outcome': 'Outcome_output1',
    #     'Response': 'Response_output1',
    #     'Reason': 'Reason_output1'
    # }

    #json_response['Data'].append(Data)

    #json_response['Data'].append(Data1)

    outputFilename = "outputExcel//"+str(f.filename) + "_output.xlsx"
    outputFilename=os.path.join("outputExcel",str(f.filename) + "_output.xlsx")
    i=0
    OutputFile=pd.DataFrame()
    while(i < json_response['Data'].__len__()):
        tempOutputFile = pd.DataFrame(json_response['Data'][i].items(), columns=['Field', 'Response'])
        OutputFile=OutputFile.append(tempOutputFile)
        i+=1

    OutputFile.to_excel(outputFilename)

    print("OutputFilename: " + outputFilename)

    json_response['Path']=outputFilename
    print(json_response)
    jsonOutput = json.dumps(json_response)
    return jsonOutput


@app.route('/router1', methods=['GET','POST','OPTIONS'])
def router1():
   print("Router")
   word='liability'
   words=['liability','liable']
   #import Processor.Processor as pp
   #duration=pp.x
   #print(request)
   print(request.form)
   f = request.files['file']
   print("Printing checkboxes")
   print(request.form.getlist('keys[]'))
   print(request.form.getlist('downloadfile'))

   print("File")
   #f.save((f.filename))
   doc=Document(f)
   print("Document")
   l = list()
   section = list()
   print("File"+str(f.filename))
   file="Output Documents\\"+str(f.filename)+".csv"
   sectionfile="OutputSection\\"+str(f.filename)+"_section.csv"
   for i, paragraphs in enumerate(doc.paragraphs):
      text = (paragraphs.text)
      text = ''.join(x for x in (text.encode('utf-8').decode('latin-1')) if x in string.printable)
      if str(text):
         l.append(str(text))
      for x in words:
        if (str(text.lower()).__contains__(x)):    #if [s for s in (text.lower()).split() if any(xs in s for xs in words)]:
            section.append(str(text))
            break;
   d = pd.DataFrame(l)
   sectiondf = pd.DataFrame(section)
   sectiondf.to_csv(sectionfile)
   d.to_csv(file)

   ClassifierResponse = AIEngineIntegrator('1079',sectionfile)
   print("Classifier response", ClassifierResponse)

   NERintegration = AIEngineIntegrator('1041',sectionfile)

   print("NER integration", NERintegration)
   DurationData=formulateResponseforDuration(ClassifierResponse,NERintegration)
   AmountData = formulateResponseforAmount(ClassifierResponse, NERintegration)

   # Data = {
   #         'Start_End':'StartEnd',
   #         'Gap_outcome': "outcome",
   #         'Response': "Response"
   #         }
   outcome,reason=formulateReason(DurationData,AmountData)

   Data = {
       'Confidence': DurationData['Confidence'],
       'Start_End': DurationData['Start_End'],
       'Gap_outcome': outcome,
       'Response': DurationData['Response'],
       'Reason': reason
   }

   outputFilename = f.filename + "_output.xlsx"

   OutputFile = pd.DataFrame(Data.items(), columns=['Field', 'Response'])
   OutputFile.to_excel(outputFilename)

   print("OutputFilename: "+outputFilename)
   outputFilepath="https://13.232.229.71/user/vaibhav.vijay.kotwal/files/CDAT%20Flask/OutputSection/"+outputFilename+"?download=1"
   print("outputFilepath: " + outputFilepath)
   Data = {
       'Confidence': DurationData['Confidence'],
       'Start_End':DurationData['Start_End'],
       'Gap_outcome':outcome ,
       'Response': DurationData['Response'],
       'Reason': reason,
       'path':outputFilepath
   }


   jsonOutput=json.dumps(Data)
   print(jsonOutput)
   return jsonOutput  #duration[0]


def AIEngineIntegrator(pipelineId,sectionfile):

    storage = StringIO()

    url = "https://aaie.accenture.com/aceapi/ace-pipeline-execution-service/execute/PipelineExecution/"
    filepath = sectionfile#'OutputSection\\liability_section.csv'
    c = pycurl.Curl()
    c.setopt(c.URL, url)
    # c.setopt(c.HTTPHEADER, ['enctype : multipart/form-data', 'Content-Type: multipart/form-data'])

    c.setopt(c.USERPWD, '%s:%s' % ('ace', 'ace'))

    c.setopt(c.CAINFO, certifi.where())

    c.setopt(c.HTTPPOST, [('file', (c.FORM_FILE, filepath,
                                    c.FORM_CONTENTTYPE, 'multipart/form-data',
                                    )), ('ownerUserName', 's.b.karuppasamy'), ('ownerUserId', '130'),
                          ('executorUserName', 's.b.karuppasamy'), ('executorUserId', '130'), ('pipelineId',pipelineId ),
                          ('getStatusWithResponse', 'true'), ('input', ''), ('bulletHandler', 'false'),
                          ('extractText', 'true')

                          ])
    # writing output to storage NER '1041'
    c.setopt(c.WRITEFUNCTION, storage.write)
    c.perform()
    c.close()
    # Required content
    content = storage.getvalue()
    print content
    content_str = content.decode()
    data = json.loads(content_str)

    return data

def resolveDuration1(duration):
    x=list()

    duration=re.sub('[(){}<>]', '', duration)
    for word in re.split(' |-',duration):
        try:
            if word.isdigit():
                x.append(word)
            else:
                x.append(wp.words_to_num(word))
        except Exception as error:
            continue
    #x=parserduration(x,duration)

    return x

def parserduration(x,duration):
    print("Parsing duration: "+duration)
    print("Parsing x "+str(x))
    period=str(x)
    if(int(period)<=12):
        duration=duration
    if duration.lower().__contains__('months') or duration.lower().__contains__('month'):
        print("months found")
        period=int(x)*30
        print("Parsed Duration: "+str((period)))
    elif duration.lower().__contains__('years') or duration.lower().__contains__('year'):
        print("years found")
        period = int(x) * 365
        print("Parsed Duration: " + str((period)))

    if(str(period)==str(x) and int(period)<=12):
        period = int(period) * 30
        print("Reparsed Duration: " + str((period)))

    return period

def compareDuration(duration):
    ConfigPath = "CDAT Config.xlsx"
    ConfigTbl = pd.read_excel(ConfigPath)
    response = "No gap"
    ConfigTbl.loc[ConfigTbl['Standards'] == 'Duration of liability']
    StandardDuration = int(ConfigTbl.loc[ConfigTbl['Standards'] == 'Duration of liability']["Values"].values[0])
    if (duration != StandardDuration):
        response = str(ConfigTbl.loc[ConfigTbl['Standards'] == 'Duration of liability']["Response"].values[0])
        reason = str(ConfigTbl.loc[ConfigTbl['Standards'] == 'Duration of liability']["Reason"].values[0])
        return "Gap Found", response,reason  #return True, response
    else:
        response = str(ConfigTbl.loc[ConfigTbl['Standards'] == 'Duration of liability']["Response"].values[0])
        reason = str(ConfigTbl.loc[ConfigTbl['Standards'] == 'Duration of liability']["PassingRemark"].values[0])
        return "No gap", response,reason #return False, response


def NoDuration():
    ConfigPath = "CDAT Config.xlsx"
    ConfigTbl = pd.read_excel(ConfigPath)
    response = str(ConfigTbl.loc[ConfigTbl['Standards'] == 'Duration of liability']["ResponseForNoValue"].values[0])
    reason=response
    return "Gap Found", response,reason

def NoAmount():
    ConfigPath = "CDAT Config.xlsx"
    ConfigTbl = pd.read_excel(ConfigPath)
    response = str(ConfigTbl.loc[ConfigTbl['Standards'] == 'liability cap amount']["ResponseForNoValue"].values[0])
    reason=response
    return "Gap Found", response,reason

def compareAmount(amount):
        print("Inside compareAmount")
        ConfigPath = "CDAT Config.xlsx"
        ConfigTbl = pd.read_excel(ConfigPath)
        response = "No gap"
        ConfigTbl.loc[ConfigTbl['Standards'] == 'liability cap amount']
        StandardAmount = int(ConfigTbl.loc[ConfigTbl['Standards'] == 'liability cap amount']["Values"].values[0])
        if (int(amount) > StandardAmount):
            response = str(ConfigTbl.loc[ConfigTbl['Standards'] == 'liability cap amount']["Response"].values[0])
            reason = str(ConfigTbl.loc[ConfigTbl['Standards'] == 'liability cap amount']["Reason"].values[0])
            return "Gap Found", response, reason  # return True, response
        else:
            response = str(ConfigTbl.loc[ConfigTbl['Standards'] == 'liability cap amount']["Response"].values[0])
            reason = str(ConfigTbl.loc[ConfigTbl['Standards'] == 'liability cap amount']["PassingRemark"].values[0])
            return "No gap", response, reason  # return False, response


def compareNoticePeriodBechmarking(duration):
    print("Inside compareNoticePeriodBechmarking")
    ConfigPath = "CDAT Config.xlsx"
    ConfigTbl = pd.read_excel(ConfigPath)
    response = "No gap"
    ConfigTbl.loc[ConfigTbl['Standards'] == 'benchmarking notice period']
    StandardDuration = int(ConfigTbl.loc[ConfigTbl['Standards'] == 'benchmarking notice period']["Values"].values[0])
    if (int(duration) > StandardDuration):
        response = str(ConfigTbl.loc[ConfigTbl['Standards'] == 'benchmarking notice period']["Response"].values[0])
        reason = str(ConfigTbl.loc[ConfigTbl['Standards'] == 'benchmarking notice period']["Reason"].values[0])
        return "Gap Found", response, reason  # return True, response
    else:
        response = str(ConfigTbl.loc[ConfigTbl['Standards'] == 'benchmarking notice period']["Response"].values[0])
        reason = str(ConfigTbl.loc[ConfigTbl['Standards'] == 'benchmarking notice period']["PassingRemark"].values[0])
        return "No gap", response, reason

def compareInitialPeriodBechmarking(duration):
    print("Inside compareInitialPeriodBechmarking")
    ConfigPath = "CDAT Config.xlsx"
    ConfigTbl = pd.read_excel(ConfigPath)
    response = "No gap"
    ConfigTbl.loc[ConfigTbl['Standards'] == 'benchmarking initial period']
    StandardDuration = int(ConfigTbl.loc[ConfigTbl['Standards'] == 'benchmarking initial period']["Values"].values[0])
    if (int(duration) > StandardDuration):
        response = str(ConfigTbl.loc[ConfigTbl['Standards'] == 'benchmarking initial period']["Response"].values[0])
        reason = str(ConfigTbl.loc[ConfigTbl['Standards'] == 'benchmarking initial period']["Reason"].values[0])
        return "Gap Found", response, reason  # return True, response
    else:
        response = str(ConfigTbl.loc[ConfigTbl['Standards'] == 'benchmarking initial period']["Response"].values[0])
        reason = str(ConfigTbl.loc[ConfigTbl['Standards'] == 'benchmarking initial period']["PassingRemark"].values[0])
        return "No gap", response, reason

def formulateResponseforNoticePeriodBechmarking(duration):
    print("Inside formulateResponseforNoticePeriodBechmarking")
    outcome, response, reason= compareNoticePeriodBechmarking(duration)
    Data = {


        'Gap_outcome': str(outcome),
        'Response': response,
        'Reason': reason
    }
    return Data


def formulateResponseforInitialPeriodBechmarking(duration):
    print("Inside formulateResponseforInitialPeriodBechmarking")
    outcome, response, reason = compareInitialPeriodBechmarking(duration)
    Data = {

        'Gap_outcome': str(outcome),
        'Response': response,
        'Reason': reason
    }
    return Data


def formulateResponseforDuration(ClassifierResponse,NERintegration):
    Confidence = (json.loads(((ClassifierResponse['finalOutPut']['finalString']))))['Confidence']
    print("Formulate Response for duration")
    Target = (json.loads(((ClassifierResponse['finalOutPut']['finalString']))))['Target']
    # duration="No duration"
    duration = "0"
    for i in range(NERintegration['finalOutPut']['finalList'].__len__()):
        if (NERintegration['finalOutPut']['finalList'][i]['nerDateOutput']):
            duration = (NERintegration['finalOutPut']['finalList'][i]['nerDateOutput'][0]['nervalue'])
            sentence = (NERintegration['finalOutPut']['finalList'][i]['sentence'])
            start = int(NERintegration['finalOutPut']['finalList'][i]['nerDateOutput'][0]['start'])
            end = int(NERintegration['finalOutPut']['finalList'][i]['nerDateOutput'][0]['end'])
            Start_End = sentence  # [start-3:end+4]
    print('Confidence ' + Confidence)
    if str(duration) is "0":
        outcome, response,reason = NoDuration()
        print("Duration not found")
        Start_End="Nothing found"
    else:
        x = resolveDuration(duration)

        # print('Duration found:' + str(x[0]))
        # outcome, response,reason = compareDuration(int(x[0]))
        print('Duration found:' + str(x))
        outcome, response,reason = compareDuration(int(x))

        print("Before curtailing",Start_End)
        Start_End = ' '.join(str(Start_End).split(',')[1:])
        print("After curtailing", Start_End)
    # Start_End=Start_End.split(',')[1:]

    print("Outcome Response", outcome, response,reason)

    #reason = 'Because of the mismatch in the duration of the liability we have categorized it as a gap'
                #'Duration': str(x[0]),
    Data = {
            'Confidence': Confidence,
            'Start_End': str(Start_End),
            'Gap_outcome': str(outcome),
            'Response': response,
            'Reason': reason
            }

    return Data


def formulateResponseforAmount(ClassifierResponse,NERintegration):
    Confidence = (json.loads(((ClassifierResponse['finalOutPut']['finalString']))))['Confidence']
    print("Formulate Response for Amount")
    Target = (json.loads(((ClassifierResponse['finalOutPut']['finalString']))))['Target']

    percent = "0"
    for i in range(NERintegration['finalOutPut']['finalList'].__len__()):
        if (NERintegration['finalOutPut']['finalList'][i]['nerPercentageOutput']):
            percent = (NERintegration['finalOutPut']['finalList'][i]['nerPercentageOutput'][0]['nervalue'])
            sentence = (NERintegration['finalOutPut']['finalList'][i]['sentence'])
            start = int(NERintegration['finalOutPut']['finalList'][i]['nerPercentageOutput'][0]['start'])
            end = int(NERintegration['finalOutPut']['finalList'][i]['nerPercentageOutput'][0]['end'])
            Start_End = sentence  # [start-3:end+4]
    print('Confidence ' + Confidence)
    if str(percent) is "0":
        outcome, response,reason = NoAmount()
        print("Amount not found")
        Start_End="We could not find liability section in the given SOW"
    else:
        percent = ''.join(e for e in percent if e.isdigit())

        print('Amount found:' + percent)
        outcome, response,reason = compareAmount(percent)
        Start_End = ' '.join(str(Start_End).split(',')[1:])
    # Start_End=Start_End.split(',')[1:]

    print("Outcome Response", outcome, response,reason)

    #reason = 'Because of the mismatch in the duration of the liability we have categorized it as a gap'
                #'Duration': str(x[0]),
    Data = {
            'Confidence': Confidence,
            'Start_End': str(Start_End),
            'Gap_outcome': str(outcome),
            'Response': response,
            'Reason': reason
            }



    return Data

def formulateReason(DurationData,AmountData):
    if (DurationData['Gap_outcome'] == "Gap Found" and  AmountData['Gap_outcome'] == "No gap"):
        outcome="Gap Found"
        reason=DurationData['Reason']
    if (DurationData['Gap_outcome'] == "No gap" and  AmountData['Gap_outcome'] == "Gap Found"):
        outcome="Gap Found"
        reason = AmountData['Reason']
    if (DurationData['Gap_outcome'] == "Gap Found" and  AmountData['Gap_outcome'] == "Gap Found"):
        outcome="Gap Found"
        reason='Because of the mismatch in the amount & duration of the liability we have categorized it as a gap'
    if (DurationData['Gap_outcome'] == "No gap" and  AmountData['Gap_outcome'] == "No gap"):
        outcome = "No Gap Found"
        reason="Liability section is present with amount and liability duration as per Accenture guidelines"

    return outcome,reason

def formulateReasonforBenchmarking(NoticePeriodData,InitialPeriodData):
    if (NoticePeriodData['Gap_outcome'] == "Gap Found" and  InitialPeriodData['Gap_outcome'] == "No gap"):
        outcome="Gap Found"
        reason=NoticePeriodData['Reason']
    if (NoticePeriodData['Gap_outcome'] == "No gap" and  InitialPeriodData['Gap_outcome'] == "Gap Found"):
        outcome="Gap Found"
        reason = InitialPeriodData['Reason']
    if (NoticePeriodData['Gap_outcome'] == "Gap Found" and  InitialPeriodData['Gap_outcome'] == "Gap Found"):
        outcome="Gap Found"
        reason='Because of the mismatch in the notice period & initial notice period of the liability we have categorized it as a gap'
    if (NoticePeriodData['Gap_outcome'] == "No gap" and  InitialPeriodData['Gap_outcome'] == "No gap"):
        outcome = "No Gap Found"
        reason="Benchmarking section is present with notice period and initial period as per Accenture guidelines"

    return outcome,reason


def createSectionFile(f,words):
    print("File")

    # f.save((f.filename))
    doc = Document(f)
    print("Document")
    l = list()
    section = list()
    print("File" + str(f.filename))
    file = "Output Documents\\" + str(f.filename) + ".csv"
    file=os.path.join("Output Documents",str(f.filename) + ".csv")
    sectionfile = "OutputSection\\" + str(f.filename) + "_section.csv"
    sectionfile=os.path.join("OutputSection",str(f.filename) + "_section.csv")
    for i, paragraphs in enumerate(doc.paragraphs):
        text = (paragraphs.text)
        text = ''.join(x for x in (text.encode('utf-8').decode('latin-1')) if x in string.printable)

        if str(text):
            l.append(str(text))
        for x in words:
            if (str(text.lower()).__contains__(
                    x)):  # if [s for s in (text.lower()).split() if any(xs in s for xs in words)]:
                text = preprocess(text)
                section.append(str(text))
                break;
    d = pd.DataFrame(l)
    sectiondf = pd.DataFrame(section)
    sectiondf.to_csv(sectionfile)
    d.to_csv(file)

    return sectionfile

def preprocess(text):
    dictword = {"twice": "200%", "thrice": "300%","exceeds":"200%"}  #,"exceed":"200%","exceeds":"200%","exceeded":"200%"
    for word in text.split():
        if (word in dictword.keys()):
            text=text.replace(word, dictword[word])

    return text

def get_answer_json(path, question):
        csvdata = pd.read_csv(
            path, header=None,
            names=['No', 'text'])
        text = str()
        for x in csvdata['text'][2:]:
            text = text + " " + (x) + " "

        #print(text)

        request_input = {
            "passage": "",
            "question": "",
            "text_type": "freetext",
            "file_path": "samples/contract.docx",
            "sessionParam": {"contextId": "24e81a62e5d045a22b0ff175d3c10e3cd273f4649f61d165a8bdb7929367a3c6",
                             "SenderId": "textanalysis", "SenderName": "textanalysis"}
        }

        request_input["passage"] = text
        request_input["question"] = question
        request_input_json = json.dumps(request_input)

        url = 'http://13.126.124.103/text-analysis-accelerator-agent/aaip/machinecomprehension/service/' #13.126.124.103 #13.126.249.116
        print("Request "+request_input_json)
        response = requests.post(url, data=request_input_json, headers={"Content-Type": "application/json"})

        print("Machine comprehension integration status "+str(response.status_code))
        #print(str(response.json()))
        #jSonOutput = response.json()

        NERvalue=str(response.json()["message"]['msg'])
        print("NER value from machine comprehension "+(NERvalue))
        start_end=""
        for sent in text.split('.'):
            if (NERvalue in sent):
                print(sent)
                start_end=sent
                break

        return NERvalue,start_end

def resolveDuration(duration):
        print ("Resolving benchmarking NER")
        x = list()
        print("Unprocessed duration "+duration)
        tempduration=duration
        duration = re.sub('[(){}<>]', '', duration)
        for word in re.split(' |-', duration):
            try:
                if word.isdigit():
                    x.append(word)
            except Exception as error:
                continue

        if(x.__len__()==0):
            num_word = [y for y in re.split(' |-', duration) if (y in wp.numwords.keys())]
            x=[wp.words_to_num("-".join(num_word))]

        period = parserduration(x[0], duration)

        print("Period value "+str(period))

        return period

def get_notice_period(path):
        question = "what is the duration of notice period of benchmarking?"
        print(question)
        notice_period,start_end = get_answer_json(path, question)
        notice_period = resolveDuration((notice_period))

        return notice_period,start_end

def get_initial_period(path):
        question = "what is the duration of initial period of benchmarking?"
        print(question)
        initial_period,start_end = get_answer_json(path, question)

        initial_period = resolveDuration((initial_period))

        return initial_period,start_end

def clearDirectory():
    path=os.getcwd()+"\\OutputFile"
    path=os.path.join(os.getcwd(),"OutputFile")
    from os import walk
    print("Cleaning "+path)
    for (filename) in os.listdir(path):
        #os.remove(path + "\\" + filename)
        #print("Removed " + filename)
        if (filename.__contains__('~$')):
            os.remove(path + "\\" + filename)
            print("--Removed-- "+filename)
    # for (filename) in os.listdir(path):
    #         os.remove(path + "\\" + filename)
    return


def get_pagenumbers(indexfilepath):

    try:
        indexfile=pd.read_csv(indexfilepath)
        indices=indexfile['1'].values
        page_nos=str(list(set(indices.flat))).strip('[]')
    except Exception as error:
        page_nos = '--'
        print("No page found")

    return page_nos



if __name__ == "__main__":
    app.run(debug=True, threaded=True, host='127.0.0.1', port=4999)

#<img src="img/load.gif"  id="loader">
# alert(obj['Start_End']);
#       $("#outputBox").show();
#       $(".outputText").text(obj['Start_End']);
#       $(".feedbackHead").text(obj['Response']);
#       $(".gap").text(obj['Gap_outcome']);
#c=' '.join(str(txt).split(',')[1:])


#for i in range(NERintegration['finalOutPut']['finalList'].__len__()):
#     if (NERintegration['finalOutPut']['finalList'][i]['nerPercentageOutput']):
#         percent = (NERintegration['finalOutPut']['finalList'][i]['nerPercentageOutput'][0]['nervalue'])
#percent=''.join(e for e in percent if e.isdigit())

#@2018 Accenture.All Rights Reserved
#http://127.0.0.1:5000/router
#http://13.232.229.71/router

#https://13.232.229.71/user/vaibhav.vijay.kotwal/files/CDAT%20Flask/OutputSection/AI%20Sample%20%20duration%20conflict.docx_section.csv?download=1
#https://13.232.229.71/user/vaibhav.vijay.kotwal/files/CDAT%20Flask/"+result1.Path+"?download=1